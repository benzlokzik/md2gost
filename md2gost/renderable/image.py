import logging
from copy import copy
from io import BytesIO
from typing import Generator
from os import environ
import os.path

import requests
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Parented, Length
from docx.text.paragraph import Paragraph

from .caption import Caption, CaptionInfo
from .renderable import Renderable
from .requires_numbering import RequiresNumbering
from ..layout_tracker import LayoutState
from ..rendered_info import RenderedInfo
from ..sub_renderable import SubRenderable
from ..util import create_element


class Image(Renderable, RequiresNumbering):
    def __init__(self, parent: Parented, path: str, caption_info: CaptionInfo | None = None):
        super().__init__("Рисунок")
        self._parent = parent
        self._caption_info = caption_info
        self._docx_paragraph = Paragraph(create_element("w:p"), parent)
        self._docx_paragraph.paragraph_format.space_before = 0
        self._docx_paragraph.paragraph_format.space_after = 0
        self._docx_paragraph.paragraph_format.first_line_indent = 0
        self._docx_paragraph.paragraph_format.line_spacing = 1
        self._docx_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._invalid = False

        run = self._docx_paragraph.add_run()

        if path.startswith("http"):
            bytesio = BytesIO()
            bytesio.write(requests.get(path).content)
            self._image = run.add_picture(bytesio)
        else:
            try:
                self._image = run.add_picture(os.path.join(environ['WORKING_DIR'], os.path.expanduser(path)))
            except FileNotFoundError:
                logging.warning(f"Invalid image path: {path}, skipping...")
                self._invalid = True

        self._number = None

    def set_number(self, number: int):
        self._number = number

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | SubRenderable, None, None]:
        if self._invalid:
            yield from []
            return

        # limit width
        if self._image.width > layout_state.max_width:
            height_by_width = self._image.height / self._image.width
            self._image.width = layout_state.max_width
            self._image.height = Length(self._image.width * height_by_width)

        # limit height
        if self._image.height > layout_state.max_height:
            width_by_height = self._image.width / self._image.height
            self._image.height = layout_state.max_height
            self._image.width = Length(self._image.height * width_by_height)

        height = self._image.height

        if layout_state.remaining_page_height < height:
            height += layout_state.remaining_page_height

        yield (rendered_image := RenderedInfo(self._docx_paragraph, Length(height)))

        layout_state.add_height(rendered_image.height)

        caption = Caption(self._parent, "Рисунок", self._caption_info, self._number, False)
        caption.center()

        yield from caption.render(rendered_image, copy(layout_state))
