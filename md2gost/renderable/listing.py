from copy import copy
import os
from typing import Generator, Callable

from docx.oxml import CT_Tbl
from docx.shared import Length, Pt, RGBColor, Twips
from docx.table import Table

from pygments import highlight
from pygments.formatter import Formatter
from pygments.lexers import get_lexer_by_name

from .caption import Caption, CaptionInfo
from .paragraph import Paragraph
from .renderable import Renderable
from .requires_numbering import RequiresNumbering
from ..docx_elements import create_table
from ..layout_tracker import LayoutState
from ..rendered_info import RenderedInfo
from ..sub_renderable import SubRenderable


class DocxParagraphPygmentsFormatter(Formatter):
    def __init__(self, paragraphs: list[Paragraph], creator: Callable[[], Paragraph], **options):
        Formatter.__init__(self, style="sas", **options)
        self._creator = creator
        self._paragraphs = paragraphs
        self._styles = {}

        for token, style in self.style:
            self._styles[token] = style

    def _add_run_to_last_paragraph(self, text, style):
        self._paragraphs[-1].add_run(text, style["bold"] or None, style["italic"] in style or None,
                                     RGBColor.from_string(style['color']) if style['color'] else None)

    def format(self, tokensource, outfile):
        self._paragraphs.append(self._creator())
        for ttype, value in tokensource:
            style = self._styles[ttype]
            lines = iter(value.split("\n"))
            self._add_run_to_last_paragraph(next(lines), style)
            for line in lines:
                self._paragraphs.append(self._creator())
                self._add_run_to_last_paragraph(line, style)
        self._paragraphs.pop(-1)  # remove last empty line


LISTING_OFFSET = Pt(31) - Twips(108 * 2)  # todo: fix


class Listing(Renderable, RequiresNumbering):
    def __init__(self, parent, language: str, caption_info: CaptionInfo):
        super().__init__("Листинг")
        self._caption_info = caption_info
        self._language = language
        self._parent = parent
        self.paragraphs: list[Paragraph] = []
        self._number = None

    def _create_table(self, parent, width: Length):
        # todo: style inheritance
        left_margin = Twips(int(
            parent.part.styles["Normal Table"]._element.xpath("w:tblPr/w:tblCellMar/w:left")[0].attrib[
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"]))
        right_margin = Twips(int(
            parent.part.styles["Normal Table"]._element.xpath("w:tblPr/w:tblCellMar/w:right")[0].attrib[
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"]))

        return create_table(parent, 1, 1, width + left_margin + right_margin)

    def set_text(self, text: str):
        def create_paragraph() -> Paragraph:
            paragraph = Paragraph(self._parent)
            paragraph.style = "Code"
            return paragraph

        text = text.removesuffix("\n")

        if self._language and "SYNTAX_HIGHLIGHTING" in os.environ and os.environ["SYNTAX_HIGHLIGHTING"] == "1":
            formatter = DocxParagraphPygmentsFormatter(self.paragraphs, lambda: create_paragraph())
            highlight(text, get_lexer_by_name(self._language), formatter)
        else:
            for line in text.removesuffix("\n").split("\n"):
                paragraph = create_paragraph()
                paragraph.add_run(line)
                self.paragraphs.append(paragraph)

    def set_number(self, number: int):
        self._number = number

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | SubRenderable, None, None]:
        caption_rendered_infos = list(
            Caption(self._parent, "Листинг", self._caption_info, self._number, True)
            .render(previous_rendered, copy(layout_state))
        )
        layout_state.add_height(sum([info.height for info in caption_rendered_infos]))
        yield from caption_rendered_infos

        table = self._create_table(self._parent, layout_state.max_width)
        previous = None

        table_height = Pt(1)  # table borders, 4 eights of point for each border

        # if first line doesn't fit move listing to the next page
        paragraph_layout_state = copy(layout_state)
        paragraph_layout_state.max_width -= LISTING_OFFSET
        paragraph_rendered_info = next(self.paragraphs[0].render(previous, paragraph_layout_state))
        if paragraph_rendered_info.height + table_height > layout_state.remaining_page_height:
            table_height += layout_state.remaining_page_height
            layout_state.add_height(layout_state.remaining_page_height)

        for paragraph in self.paragraphs:
            paragraph_layout_state = copy(layout_state)
            paragraph_layout_state.max_width -= LISTING_OFFSET
            paragraph_rendered_info = next(paragraph.render(previous, paragraph_layout_state))

            if paragraph_rendered_info.height > layout_state.remaining_page_height:  # todo add before after
                table_rendered_info = RenderedInfo(table, table_height)
                yield table_rendered_info

                table_height = Pt(1)  # table borders, 4 eights of point for each border

                continuation_paragraph = Paragraph(self._parent)
                continuation_paragraph.add_run(f"Продолжение листинга {self._number}")
                continuation_paragraph.style = "Caption"
                continuation_paragraph.first_line_indent = 0
                continuation_paragraph.page_break_before = True

                continuation_rendered_info = next(
                    continuation_paragraph.render(None, copy(layout_state)))

                layout_state.add_height(continuation_rendered_info.height)
                yield continuation_rendered_info

                table = self._create_table(self._parent, layout_state.max_width)

                previous = None

                paragraph_layout_state = copy(layout_state)
                paragraph_layout_state.max_width -= LISTING_OFFSET
                paragraph_rendered_info = next(paragraph.render(previous, paragraph_layout_state))

            table._cells[0]._element.append(paragraph_rendered_info.docx_element._element)
            layout_state.add_height(paragraph_rendered_info.height)
            table_height += paragraph_rendered_info.height

            previous = paragraph_rendered_info

        yield RenderedInfo(table, table_height)
