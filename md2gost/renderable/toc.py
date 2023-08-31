from copy import copy
from typing import Generator

from docx.enum.text import WD_TAB_LEADER, WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Parented, Pt
from docx.text.run import Run

from . import Paragraph
from .page_break import PageBreak
from .renderable import Renderable
from ..layout_tracker import LayoutState
from ..rendered_info import RenderedInfo
from ..util import create_element


def create_field(parent: Parented, text: str, instr_text: str):
    run = Run(create_element("w:r"), None)
    run._t = create_element("w:t")

    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "begin"
    }))
    run._element.append(create_element("w:instrText", {
        "xml:space": "preserve"
    }, instr_text))
    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "separate"
    }))
    run._element.append(run._t)
    if text is not None:
        run._t.text = text
    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "end"
    }))
    return run


class ToC(Renderable):
    """
    Items are added by calling add_item(level, title, page) method.
    After the document is fully rendered fill must be called.
    """

    def __init__(self, parent: Parented):
        self._parent = parent
        self._paragraph = Paragraph(parent)
        self._paragraph._docx_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self._paragraph.first_line_indent = 0
        self._items: list[tuple[int, str, int]] = []
        pass

    def add_item(self, level: int, title: str, page: int):
        self._items.append((level, title, page))

    def fill(self):
        p = self._paragraph._docx_paragraph
        p.paragraph_format.tab_stops.add_tab_stop(
            p.part.document.sections[0].page_width-p.part.document.sections[0].left_margin-p.part.document.sections[0].right_margin,
            alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        p.paragraph_format.tab_stops.add_tab_stop(0, alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)

        numbering = [0 for _ in range(10)]
        for level, title, page in self._items:
            numbering[level-1] += 1
            for i in range(level, len(numbering)):
                numbering[i] = 0
            p.add_run("    "*(level-1))
            p.add_run(".".join([str(x) for x in numbering[:level]])+". ")
            p.add_run(title)
            p.add_run(f"\t{page}")
            p.add_run("\n")

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | Renderable, None, None]:
        for rendered_info in self._paragraph.render(previous_rendered, copy(layout_state)):
            yield RenderedInfo(rendered_info.docx_element, 0)
        yield from PageBreak(self._parent).render(None, copy(layout_state))
