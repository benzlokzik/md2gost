from copy import copy
from typing import Generator, Any

from docx.shared import Length, Parented, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.paragraph import Run as DocxRun
from docx.enum.text import WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE

from . import Renderable
from .image import Image
from .paragraph_sizer import ParagraphSizer
from ..layout_tracker import LayoutState
from ..sub_renderable import SubRenderable
from ..util import create_element
from ..rendered_info import RenderedInfo
from ..latex_math import latex_to_omml, inline_omml


class Link:
    def __init__(self, url, docx_paragraph: DocxParagraph):
        self._docx_paragraph = docx_paragraph
        r_id = docx_paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        self._hyperlink = create_element("w:hyperlink", {
            "r:id": r_id
        })

    def add_run(self, text: str, is_bold: bool = None, is_italic: bool = None, color: RGBColor = None,
                    strike_through: bool = None):

        parts = text.split("-")
        for i, part in enumerate(parts):
            docx_run = DocxRun(create_element("w:r"), self._docx_paragraph)
            self._hyperlink.append(docx_run._element)
            docx_run.text = text
            docx_run.style = "Hyperlink"
            docx_run.bold = is_bold
            docx_run.italic = is_italic
            docx_run.font.color.rgb = color
            docx_run.font.strike = strike_through
            if i != len(parts) - 1:
                self._hyperlink.append(create_element("w:r", [create_element("w:noBreakHyphen")]))

    @property
    def element(self):
        return self._hyperlink


class Paragraph(Renderable):
    def __init__(self, parent: Parented):
        self._parent = parent
        self._docx_paragraph = DocxParagraph(create_element("w:p"), parent)
        self._docx_paragraph.style = "Normal"
        self._images: list[Image] = []

    def add_run(self, text: str, is_bold: bool = None, is_italic: bool = None, color: RGBColor = None,
                strike_through: bool = None):
        # replace all hyphens with non-breaking hyphens
        parts = text.split("-")
        for i, part in enumerate(parts):
            docx_run = self._docx_paragraph.add_run(part)
            docx_run.bold = is_bold
            docx_run.italic = is_italic
            docx_run.font.color.rgb = color
            docx_run.font.strike = strike_through
            if i != len(parts)-1:
                self._docx_paragraph.add_run()._element.\
                    append(create_element("w:noBreakHyphen"))

    def add_image(self, path: str):
        self._images.append(Image(self._parent, path))

    def add_link(self, url: str):
        link = Link(url, self._docx_paragraph)
        self._docx_paragraph._p.append(link.element)
        return link

    def add_inline_equation(self, formula: str):
        # omml = inline_omml(latex_to_omml(formula))
        # for r in omml.xpath("//m:r", namespaces=omml.nsmap):
        #     r.append(create_element("w:rPr", [
        #         create_element("w:sz", {"w:val": "24"}),
        #         create_element("w:szCs", {"w:val": "24"}),
        #     ]))
        # self._docx_paragraph._element.append(omml)
        self.add_run(formula, is_italic=True)

    @property
    def page_break_before(self) -> bool:
        return self._docx_paragraph.paragraph_format.page_break_before

    @page_break_before.setter
    def page_break_before(self, value: bool):
        self._docx_paragraph.paragraph_format.page_break_before = value

    @property
    def style(self):
        return self._docx_paragraph.style

    @style.setter
    def style(self, value: str):
        self._docx_paragraph.style = value

    @property
    def first_line_indent(self):
        return self._docx_paragraph.paragraph_format.first_line_indent

    @first_line_indent.setter
    def first_line_indent(self, value: Length):
        self._docx_paragraph.paragraph_format.first_line_indent = value

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | SubRenderable, None, None]:
        remaining_space = layout_state.remaining_page_height

        if self.page_break_before:
            layout_state.add_height(layout_state.remaining_page_height)
        if self._docx_paragraph.text or not self._images:
            height_data = ParagraphSizer(
                self._docx_paragraph,
                previous_rendered.docx_element
                          if previous_rendered and isinstance(previous_rendered.docx_element, DocxParagraph) else None,
                          layout_state.max_width).calculate_height()

            if layout_state.current_page_height == 0 and layout_state.page > 1:
                height_data.before = 0

            fitting_lines = 0
            for lines in range(1, height_data.lines+1):
                if height_data.before + ((lines - 1) * height_data.line_spacing + 1) * height_data.line_height \
                        > layout_state.remaining_page_height:
                    break
                fitting_lines += 1

            if fitting_lines == height_data.lines:
                # the whole paragraph fits page
                height = min(height_data.full, layout_state.remaining_page_height)
            elif fitting_lines <= 1 or (height_data.lines-fitting_lines == 1 and height_data.lines == 3):
                # if only no or only one line fits the page, paragraph goes to the next page
                height = layout_state.remaining_page_height + height_data.full
            elif height_data.lines-fitting_lines == 1:
                # if all lines except last fit the page, the last two lines go to the new page
                height = layout_state.remaining_page_height + \
                         height_data.before + height_data.line_height * height_data.line_spacing * 2 \
                         + height_data.after
            else:
                height = layout_state.remaining_page_height + \
                         height_data.before + height_data.line_height * height_data.line_spacing * \
                         (height_data.lines-fitting_lines) + height_data.after

            if self.page_break_before:
                height += remaining_space

            yield (previous_rendered := RenderedInfo(self._docx_paragraph, Length(height)))
            layout_state.add_height(height)

        images = iter(self._images)

        for image in images:
            rendered_image = list(image.render(previous_rendered, copy(layout_state)))
            rendered_image_height = sum([x.height for x in rendered_image])
            if rendered_image:
                previous_rendered = rendered_image[-1]
            if rendered_image_height <= layout_state.remaining_page_height:
                yield SubRenderable(image, False)
                layout_state.add_height(rendered_image_height)
            else:
                yield SubRenderable(image, True)
                break

        yield from images
