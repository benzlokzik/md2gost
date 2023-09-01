from copy import copy
from dataclasses import dataclass
from typing import Generator

from docx.shared import Parented
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from md2gost.layout_tracker import LayoutState
from md2gost.renderable import Renderable
from md2gost.rendered_info import RenderedInfo
from .paragraph_sizer import ParagraphSizer
from .requires_numbering import RequiresNumbering
from ..util import create_element


@dataclass
class CaptionInfo:
    unique_name: str
    text: str | None


class Caption(Renderable):
    def __init__(self, parent: Parented, category: str, caption_info: CaptionInfo | None,
                 number: int = None, before=True):
        self._parent = parent
        self._before = before
        self._docx_paragraph = DocxParagraph(create_element("w:p"), parent)

        self._docx_paragraph.style = "Caption"
        self._docx_paragraph.add_run(f"{category} ")
        self._numbering_run = self._docx_paragraph.add_run(str(number) if number else "?")
        if caption_info and caption_info.text:
            self._docx_paragraph.add_run(f" - {caption_info.text}")

    def center(self):
        self._docx_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState) -> Generator[
        "RenderedInfo | Renderable", None, None]:
        height_data = ParagraphSizer(
            self._docx_paragraph,
            previous_rendered.docx_element
            if previous_rendered and isinstance(previous_rendered.docx_element, DocxParagraph) else None,
            layout_state.max_width
        ).calculate_height()

        # if three more lines don't fit, move it to the next page (so there is no only caption on the end of the page)
        if self._before and ((height_data.lines + 2 - 1) * height_data.line_spacing + 1) * height_data.line_height \
                > layout_state.remaining_page_height:
            self._docx_paragraph.paragraph_format.page_break_before = True
            height_data = ParagraphSizer(
                self._docx_paragraph,
                None,
                layout_state.max_width
            ).calculate_height()

        yield RenderedInfo(self._docx_paragraph, height_data.full + (layout_state.remaining_page_height
                                                                     if self._docx_paragraph.paragraph_format.page_break_before else 0))
