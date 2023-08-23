import logging
from collections import defaultdict
from io import BytesIO

from PIL import Image
from PIL.ImageDraw import ImageDraw
from docx.document import Document
from docx.shared import Parented, Length, Pt, Cm

from docx.oxml import parse_xml, register_element_cls, CT_P
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.text.paragraph import Paragraph

from md2gost.util import create_element
from .renderer import BOTTOM_MARGIN

EMUS_PER_PX = Pt(1) * 72/96


# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)

        try:
            pic.nvPicPr.xpath("pic:cNvPicPr")[0].append(
                create_element("a:picLocks", {
                    "noChangeAspect": "1",
                    "noMove": "1",
                    "noResize": "1",
                    "noRot": "1",
                })
            )
        except IndexError:
            pass

        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
                '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
                '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
                '           %s>\n'
                '  <wp:simplePos x="0" y="0"/>\n'
                '  <wp:positionH relativeFrom="page">\n'
                '    <wp:posOffset>%d</wp:posOffset>\n'
                '  </wp:positionH>\n'
                '  <wp:positionV relativeFrom="page">\n'
                '    <wp:posOffset>%d</wp:posOffset>\n'
                '  </wp:positionV>\n'
                '  <wp:extent cx="914400" cy="914400"/>\n'
                '  <wp:wrapNone/>\n'
                '  <wp:docPr id="666" name="unnamed"/>\n'
                '  <wp:cNvGraphicFramePr>\n'
                '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
                '  </wp:cNvGraphicFramePr>\n'
                '  <a:graphic>\n'
                '    <a:graphicData uri="URI not set"/>\n'
                '  </a:graphic>\n'
                '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)


# refer to docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)

def to_px(length: Length) -> int:
    return round(length / EMUS_PER_PX)


class _Page:
    def __init__(self, width: Length, height: Length,
                 margin_left: Length, margin_top: Length,
                 margin_right: Length, margin_bottom: Length):
        self._left_offset = margin_left
        self._right_offset = margin_right
        self._max_height = height - margin_bottom

        self._width, self._height = width, height
        self._image = Image.new("RGBA", (to_px(width), to_px(height)))
        self._draw = ImageDraw(self._image)
        self._offset = Length(margin_top)

        self._colors = [
            (255, 0, 0, 100),  # red
            (0, 255, 0, 100),  # green
            (0, 0, 255, 100),  # blue
        ]
        self._i = 0

    @classmethod
    def from_document(cls, document: Document):
        return cls(
            document.sections[0].page_width,
            document.sections[0].page_height,
            document.sections[0].left_margin,
            document.sections[0].top_margin,
            document.sections[0].right_margin,
            BOTTOM_MARGIN  # todo: fix detection
        )

    def add_height(self, height: Length) -> Length:
        """Return remaining height, that didn't fit to the page"""
        fitting = min(height, self._max_height - self._offset)
        remaining = height - fitting

        self._draw.rectangle(
            (
                (to_px(self._left_offset), to_px(self._offset)),
                (to_px(self._width-self._right_offset), to_px(self._offset + fitting))
            ),
            self._get_color()
        )
        self._offset += height
        return remaining

    @property
    def image(self) -> BytesIO:
        io = BytesIO()
        self._image.save(io, Image.registered_extensions()[".png"])
        return io

    def _get_color(self):
        color = self._colors[self._i % len(self._colors)]
        self._i += 1
        return color


class Debugger:
    def __init__(self, document: Document):
        self._document = document

        self._pages: list[_Page] = [_Page.from_document(document)]
        self._paragraphs_by_page: defaultdict[int, Paragraph] = defaultdict(lambda: None)

    def add(self, docx_element: Parented, height: Length):
        remaining_height = self._current_page.add_height(height)

        if not self._paragraphs_by_page[len(self._pages)-1] and isinstance(docx_element, Paragraph):
            self._paragraphs_by_page[len(self._pages)-1] = docx_element

        while remaining_height:
            self._pages.append(_Page.from_document(self._document))
            remaining_height = self._current_page.add_height(remaining_height)

    def after_rendered(self):
        """Must be called after rendering is finished"""
        if not self._document.paragraphs:
            logging.debug("No paragraphs found, can't add debug info")
            return

        for i in range(len(self._pages)):
            if not self._paragraphs_by_page[i]:
                logging.debug(f"Skipping page {i} as there are no paragraphs")
                continue
            add_float_picture(
                self._paragraphs_by_page[i],
                self._pages[i].image,
                self._document.sections[0].page_width
            )

    @property
    def _current_page(self):
        return self._pages[-1]
